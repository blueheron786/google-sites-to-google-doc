import re
from collections import deque
from typing import List, Tuple, Set
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from requests_html import HTMLSession
from urllib.parse import urljoin

# Replace with your Google Sites suffix: sites.google.com/site/YOUR_SITE_SUFFIX_HERE
SITE_SUFFIX = "YOUR_SITE_SUFFIX_HERE"
OUTPUT_FILE_NAME = f"{SITE_SUFFIX}.docx"
BASE_URL = f"https://sites.google.com/site/{SITE_SUFFIX}"

session = HTMLSession()
visited: Set[str] = set()
to_visit = deque([BASE_URL])

# Initialize document
doc = Document()
body = doc.element.body

# Add Heading1 style (20pt)
styles = doc.styles
heading_style = doc.styles['Heading 1']
heading_style.font.size = Pt(20)
heading_style.paragraph_format.keep_together = True
heading_style.paragraph_format.keep_with_next = True
heading_style.paragraph_format.space_before = Pt(24)  # 480 twips = 24pt
heading_style.paragraph_format.space_after = Pt(0)


# Add numbering for bullets (simplified approach)
doc.add_paragraph("", style="List Bullet")  # This triggers bullet style creation

def normalize_url(href: str) -> str:
    """Normalize URLs to absolute format."""
    return urljoin("https://sites.google.com/", href).replace("//", "/").replace("https:/", "https://")

def clean_title(title: str) -> str:
    """Clean the title for heading use."""
    # Remove or replace problematic characters for Google Docs anchors
    cleaned = re.sub(r"[^\w\s\-]", "", title)  # Keep letters, numbers, spaces, dashes
    return "Untitled" if not cleaned.strip() else cleaned

# Collect page data
pages: List[Tuple[str, str, List[Tuple[str, str]]]] = []

print("🕷️ Starting web crawl...")

while to_visit:
    url = to_visit.popleft()
    if url in visited:
        continue

    visited.add(url)

    try:
        response = session.get(url)
        response.raise_for_status()
        
        # Enqueue more discovered internal links
        links = [normalize_url(a.attrs["href"]) for a in response.html.find("a[href]") 
                if a.attrs.get("href", "").startswith(f"/site/{SITE_SUFFIX}/")]
        
        for link in links:
            if link not in visited:
                to_visit.append(link)

        # Remove header if it exists
        header = response.html.find("header", first=True)
        if header:
            header.element.drop_tree()

        content = response.html.find("body", first=True)
        if not content:
            print(f"⚠️ Skipping page {url}, can't find anything in main content.")
            continue

        # Extract paragraphs and list items
        blocks: List[Tuple[str, str]] = []
        
        for element in content.element.iter():
            tag = element.tag.lower()
            text = element.text_content().strip() if element.text_content() else ""
            
            if tag == "p" and not any(p.tag.lower() == "li" for p in element.iterancestors()):
                if text:
                    blocks.append(("p", text))
            elif tag == "li":
                if text:
                    blocks.append(("li", text))

        if not blocks:
            print(f"⚠️ Skipping page {url}, no content found.")
            continue

        title = (response.html.find("h1", first=True) or {}).text or url
        title = title.strip()
        
        pages.append((url, title, blocks))
        print("🔹", end="", flush=True)  # Progress indicator

    except Exception as ex:
        print(f"❌ Failed to fetch {url}: {str(ex)}")

# Sort pages alphabetically by title
pages.sort(key=lambda x: x[1])

print("\n✍️ Writing to document...")

for i, (url, title, blocks) in enumerate(pages):
    # Add heading
    doc.add_paragraph(clean_title(title), style="Heading 1")
    
    # Add content
    for block_type, text in blocks:
        if block_type == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    
    # Add page break if not last page
    if i < len(pages) - 1:
        doc.add_page_break()

doc.save(OUTPUT_FILE_NAME)
print(f"✅ Saved {len(pages)} pages to {OUTPUT_FILE_NAME}")